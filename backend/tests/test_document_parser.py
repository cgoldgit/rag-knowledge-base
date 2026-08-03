"""考题：文档解析（txt/md/pdf/docx/xlsx 提取文字）"""
import io

from openpyxl import Workbook
from pypdf import PdfWriter
from docx import Document as DocxDocument

from app.services.document_parser import UnsupportedFormatError, extract_text


class TestTxt:
    def test_txt_正常提取(self):
        content = "这是一段中文文本\n第二行"
        result = extract_text("notes.txt", io.BytesIO(content.encode("utf-8")))
        assert result == content

    def test_txt_首尾空白被清理(self):
        result = extract_text("a.txt", io.BytesIO(b"  hello world  \n"))
        assert result == "hello world"

    def test_txt_空文件返回空串(self):
        assert extract_text("empty.txt", io.BytesIO(b"")) == ""

    def test_md_正常提取(self):
        result = extract_text("doc.md", io.BytesIO("# 标题\n正文".encode("utf-8")))
        assert "# 标题\n正文" in result


class TestPdf:
    def test_pdf_逐页提取(self):
        buf = io.BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        writer.add_blank_page(width=200, height=200)
        writer.write(buf)
        buf.seek(0)
        # 空白页无文字，重点验证不报错且能处理多页
        assert extract_text("book.pdf", buf) == ""


class TestDocx:
    def test_docx_段落与表格提取(self):
        buf = io.BytesIO()
        doc = DocxDocument()
        doc.add_paragraph("第一段文字")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "型号"
        table.rows[0].cells[1].text = "X100"
        doc.save(buf)
        buf.seek(0)
        result = extract_text("spec.docx", buf)
        assert "第一段文字" in result
        assert "型号 | X100" in result

    def test_docx_空段落被跳过(self):
        buf = io.BytesIO()
        doc = DocxDocument()
        doc.add_paragraph("")
        doc.add_paragraph("有效内容")
        doc.save(buf)
        buf.seek(0)
        result = extract_text("d.docx", buf)
        assert "有效内容" in result


class TestXlsx:
    def test_xlsx_工作表与单元格提取(self):
        buf = io.BytesIO()
        wb = Workbook()
        ws = wb.active
        ws.title = "商品"
        ws.append(["名称", "价格"])
        ws.append(["手机", 2999])
        wb.save(buf)
        buf.seek(0)
        result = extract_text("goods.xlsx", buf)
        assert "[工作表: 商品]" in result
        assert "名称 | 价格" in result
        assert "手机 | 2999" in result

    def test_xlsx_空行被跳过(self):
        buf = io.BytesIO()
        wb = Workbook()
        ws = wb.active
        ws.append([None, None])
        ws.append(["有值", 1])
        wb.save(buf)
        buf.seek(0)
        result = extract_text("g.xlsx", buf)
        assert "有值 | 1" in result


class TestUnsupported:
    def test_无扩展名报错(self):
        try:
            extract_text("README", io.BytesIO(b"hi"))
        except UnsupportedFormatError:
            pass
        else:
            raise AssertionError("应抛 UnsupportedFormatError")

    def test_不支持格式报错(self):
        try:
            extract_text("photo.png", io.BytesIO(b"\x89PNG"))
        except UnsupportedFormatError:
            pass
        else:
            raise AssertionError("应抛 UnsupportedFormatError")

    def test_错误消息包含格式名(self):
        try:
            extract_text("a.zip", io.BytesIO(b"x"))
        except UnsupportedFormatError as e:
            assert "zip" in str(e)
        else:
            raise AssertionError("应抛 UnsupportedFormatError")


class TestExt:
    def test_大写扩展名同样识别(self):
        result = extract_text("NOTES.TXT", io.BytesIO("你好".encode("utf-8")))
        assert result == "你好"

    def test_markdown扩展名识别(self):
        result = extract_text("doc.markdown", io.BytesIO("# 标题".encode("utf-8")))
        assert result == "# 标题"

    def test_txt含非法UTF8字节不报错(self):
        # 二进制内容（非 UTF-8 字节）应被忽略，不抛异常
        result = extract_text("bin.txt", io.BytesIO(b"hello \xff\xfe world"))
        assert "hello" in result
        assert "world" in result
