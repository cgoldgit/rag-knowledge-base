"""文档解析服务：把各种格式的文档提取成纯文本"""
import io
from typing import BinaryIO

import pypdf
from docx import Document as DocxDocument
from openpyxl import load_workbook


class UnsupportedFormatError(Exception):
    pass


def extract_text(filename: str, file_obj: BinaryIO) -> str:
    """按扩展名解析文档，返回纯文本内容"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    text = ""

    if ext == "pdf":
        text = _extract_pdf(file_obj)
    elif ext == "docx":
        text = _extract_docx(file_obj)
    elif ext in ("txt", "md", "markdown"):
        text = file_obj.read().decode("utf-8", errors="ignore")
    elif ext == "xlsx":
        text = _extract_xlsx(file_obj)
    else:
        raise UnsupportedFormatError(f"不支持的文件格式: {ext}")

    return text.strip()


def _extract_pdf(file_obj: BinaryIO) -> str:
    """提取 PDF 文本（逐页拼接）"""
    reader = pypdf.PdfReader(file_obj)
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _extract_docx(file_obj: BinaryIO) -> str:
    """提取 Word 文本（段落 + 表格）"""
    doc = DocxDocument(file_obj)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(file_obj: BinaryIO) -> str:
    """提取 Excel 内容（每个工作表 → 逐行拼接）"""
    wb = load_workbook(file_obj, read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[工作表: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)
